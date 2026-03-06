"""
File processing tools for Nexus ADK Platform.
"""
import io
import json
import logging
import os
from typing import Any, Dict, List, Optional, Union
from pathlib import Path

import pandas as pd
from PIL import Image
import pypdf
from docx import Document

from backend.utils import validate_file_type, encode_image_to_base64

logger = logging.getLogger(__name__)

ALLOWED_IMAGE_EXTENSIONS = ['jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp']
ALLOWED_DOCUMENT_EXTENSIONS = ['pdf', 'docx', 'txt', 'md', 'csv']
ALLOWED_DATA_EXTENSIONS = ['csv', 'xlsx', 'xls', 'json']


class FileProcessingTool:
    """Tool for processing uploaded files."""

    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def process_image(
        self,
        file_data: bytes,
        filename: str,
        analyze: bool = True
    ) -> Dict[str, Any]:
        """
        Process an image file.

        Args:
            file_data: Raw image bytes
            filename: Name of the file
            analyze: Whether to perform analysis on the image

        Returns:
            Dictionary with image information and analysis
        """
        try:
            image = Image.open(io.BytesIO(file_data))

            result = {
                "success": True,
                "filename": filename,
                "format": image.format,
                "mode": image.mode,
                "size": {
                    "width": image.width,
                    "height": image.height
                },
                "info": {
                    "format": image.format,
                    "size": image.size,
                    "mode": image.mode
                }
            }

            # Get EXIF data if available
            if hasattr(image, '_getexif') and image._getexif():
                exif = image._getexif()
                if exif:
                    result["exif"] = {k: str(v) for k, v in exif.items() if isinstance(v, (str, int, float))}

            # Encode to base64 for multimodal processing
            result["base64"] = encode_image_to_base64(file_data)

            return result

        except Exception as e:
            logger.error(f"Image processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    def process_pdf(
        self,
        file_data: bytes,
        filename: str,
        extract_text: bool = True,
        max_pages: int = 50
    ) -> Dict[str, Any]:
        """
        Process a PDF file.

        Args:
            file_data: Raw PDF bytes
            filename: Name of the file
            extract_text: Whether to extract text content
            max_pages: Maximum number of pages to process

        Returns:
            Dictionary with PDF information and extracted text
        """
        try:
            pdf_file = io.BytesIO(file_data)
            reader = pypdf.PdfReader(pdf_file)

            result = {
                "success": True,
                "filename": filename,
                "num_pages": len(reader.pages),
                "metadata": reader.metadata,
                "pages": []
            }

            if extract_text:
                pages_to_process = min(len(reader.pages), max_pages)
                for i in range(pages_to_process):
                    page = reader.pages[i]
                    text = page.extract_text()
                    result["pages"].append({
                        "page_number": i + 1,
                        "text": text,
                        "char_count": len(text)
                    })

                if len(reader.pages) > max_pages:
                    result["warning"] = f"Only first {max_pages} pages processed. Total pages: {len(reader.pages)}"

            return result

        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    def process_document(
        self,
        file_data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Process a Word document.

        Args:
            file_data: Raw document bytes
            filename: Name of the file

        Returns:
            Dictionary with document content
        """
        try:
            doc_file = io.BytesIO(file_data)
            doc = Document(doc_file)

            result = {
                "success": True,
                "filename": filename,
                "paragraphs": len(doc.paragraphs),
                "content": []
            }

            for para in doc.paragraphs:
                if para.text.strip():
                    result["content"].append({
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })

            # Extract tables
            if len(doc.tables) > 0:
                result["tables"] = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    result["tables"].append(table_data)

            return result

        except Exception as e:
            logger.error(f"Document processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    def process_csv(
        self,
        file_data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Process a CSV file.

        Args:
            file_data: Raw CSV bytes
            filename: Name of the file

        Returns:
            Dictionary with CSV data and metadata
        """
        try:
            text_data = file_data.decode('utf-8')
            df = pd.read_csv(io.StringIO(text_data))

            result = {
                "success": True,
                "filename": filename,
                "shape": df.shape,
                "columns": list(df.columns),
                "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
                "preview": df.head(10).to_dict(orient='records'),
                "summary": {
                    "numeric_cols": list(df.select_dtypes(include=['number']).columns),
                    "text_cols": list(df.select_dtypes(include=['object']).columns),
                    "null_counts": df.isnull().sum().to_dict()
                }
            }

            return result

        except Exception as e:
            logger.error(f"CSV processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    def process_file(
        self,
        file_data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Process any supported file type.

        Args:
            file_data: Raw file bytes
            filename: Name of the file

        Returns:
            Processing result
        """
        ext = filename.split('.')[-1].lower()

        if ext in ALLOWED_IMAGE_EXTENSIONS:
            return self.process_image(file_data, filename)
        elif ext == 'pdf':
            return self.process_pdf(file_data, filename)
        elif ext == 'docx':
            return self.process_document(file_data, filename)
        elif ext in ['csv', 'txt', 'md']:
            return self.process_csv(file_data, filename)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {ext}",
                "filename": filename
            }


# Global file processing tool instance
file_tool = FileProcessingTool()


def process_uploaded_file(file_content: str, filename: str) -> str:
    """
    ADK Tool: Process an uploaded file and extract content.

    Args:
        file_content: Base64 encoded file content
        filename: Name of the uploaded file

    Returns:
        JSON string with processed file content
    """
    import base64
    try:
        file_data = base64.b64decode(file_content)
        result = file_tool.process_file(file_data, filename)
        return json.dumps(result, default=str)
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


def extract_image_content(file_content: str, filename: str) -> str:
    """
    ADK Tool: Extract content from an image file for vision analysis.

    Args:
        file_content: Base64 encoded image
        filename: Name of the image file

    Returns:
        JSON string with image data
    """
    import base64
    try:
        file_data = base64.b64decode(file_content)
        result = file_tool.process_image(file_data, filename)
        return json.dumps(result, default=str)
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


def extract_pdf_content(file_content: str, filename: str, max_pages: int = 50) -> str:
    """
    ADK Tool: Extract text content from a PDF file.

    Args:
        file_content: Base64 encoded PDF
        filename: Name of the PDF file
        max_pages: Maximum pages to process

    Returns:
        JSON string with extracted text
    """
    import base64
    try:
        file_data = base64.b64decode(file_content)
        result = file_tool.process_pdf(file_data, filename, max_pages=max_pages)
        return json.dumps(result, default=str)
    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})
